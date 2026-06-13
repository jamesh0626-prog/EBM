"""
AI 출력을 PDF 워크시트와 동일한 표 구조의 DOCX로 변환
"""
import io, re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


# ─── 공통 유틸 ───────────────────────────────────────────────────

def _borders(table, sz="4", color="000000"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _font(run, size=10.5, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = '맑은 고딕'
    try:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    except Exception:
        pass
    if color:
        run.font.color.rgb = RGBColor(*color)


def _cell(cell, text, bold=False, size=10.5, center=False, color=None, small=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text) if text else "")
    _font(run, size=9 if small else size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _cell_multiline(cell, lines, size=10.5, center=False):
    """셀에 여러 줄 텍스트 삽입"""
    cell.text = ""
    for i, (text, bold, color) in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(str(text) if text else "")
        _font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _heading(doc, text, size=13, bold=True, center=True, sb=10, sa=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    run = p.add_run(text)
    _font(run, size=size, bold=bold)
    return p


def _para(doc, text, size=10.5, bold=False, indent=0, sb=2, sa=3, center=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    run = p.add_run(str(text) if text else "")
    _font(run, size=size, bold=bold)
    return p


def _page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


def _set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)


# ─── AI 출력 파서 ────────────────────────────────────────────────

def _clean(text) -> str:
    """마크다운 bold/italic 제거, LaTeX 수식은 원본 유지"""
    if not text:
        return ""
    text = str(text)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text.strip()


def _get(text, pattern) -> str:
    if not text:
        return ""
    m = re.search(rf'\*\s*\*\*{pattern}[:\*]*\*\*\s*(.+?)(?=\n\*|\n##|\n###|\Z)', text, re.DOTALL|re.IGNORECASE)
    if m and m.group(1):
        return _clean(m.group(1))
    m2 = re.search(rf'\*\s*{pattern}[:\s]+(.+?)(?=\n\*|\n##|\n###|\Z)', text, re.DOTALL|re.IGNORECASE)
    if m2 and m2.group(1):
        return _clean(m2.group(1))
    return ""


def _sec(text, header) -> str:
    if not text:
        return ""
    m = re.search(rf'#{1,3}\s*{re.escape(header)}.*?\n(.*?)(?=\n#{1,3}|\Z)', text, re.DOTALL|re.IGNORECASE)
    return m.group(1).strip() if m else ""


def _iv(text, pattern):
    """Internal Validity 답변과 근거 추출"""
    m = re.search(
        rf'\*\s*\*\*.*?{pattern}.*?\*\*\s*(\[예\]|\[아니오\]|예|아니오).*?-\s*근거:\s*(.+?)(?=\n\*|\n##|\n###|\Z)',
        text, re.DOTALL|re.IGNORECASE
    )
    if m:
        ans = m.group(1).replace('[','').replace(']','').strip()
        ev  = _clean(m.group(2) or "")
        return ans, ev
    return "", ""


# ─── DOCX 생성 ──────────────────────────────────────────────────

def generate_docx(ai_text: str) -> bytes:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(3.0)

    pico_sec = _sec(ai_text, "과제 1")
    kw_sec   = _sec(ai_text, "과제 2")

    # ══════════════════════════════════════
    # 1페이지: 과제 1, 2, 3 개요
    # ══════════════════════════════════════
    _heading(doc, "년,          학년,              조", size=12, bold=False, center=True, sb=0, sa=20)

    # ── 과제 1: PICO 표 ──
    _heading(doc, "과제 1. 답변 가능한 임상질문 만들기", size=13, bold=True, center=True, sb=4, sa=6)
    t1 = doc.add_table(rows=4, cols=2)
    _borders(t1)
    pico_items = [
        ('P 환자와 문제', r'P\s*\(Patient.*?\)'),
        ('I 중재(치료)',  r'I\s*\(Intervention\)'),
        ('C 비교 치료',  r'C\s*\(Comparison\)'),
        ('O 최종 결과',  r'O\s*\(Outcome\)'),
    ]
    for i, (label, pat) in enumerate(pico_items):
        val = _get(pico_sec, pat) or _get(ai_text, pat)
        _cell(t1.rows[i].cells[0], label, size=10.5)
        _cell(t1.rows[i].cells[1], val,   size=10.5)
        _set_row_height(t1.rows[i], 1.0)
        t1.rows[i].cells[0].width = Cm(3.5)
        t1.rows[i].cells[1].width = Cm(11.5)

    doc.add_paragraph()

    # ── 과제 2: 검색어 ──
    _heading(doc, "과제 2. 검색어 선정", size=13, bold=True, center=True, sb=10, sa=6)
    kw_raw = _get(kw_sec, "Keywords") or _get(ai_text, "Keywords")
    keywords = [k.strip() for k in re.split(r'[,;，\n]', kw_raw) if k.strip()][:5]
    for i in range(5):
        kw = keywords[i] if i < len(keywords) else ""
        p = _para(doc, f"{i+1}.  {kw}", sa=5)
        if not kw:
            run = p.runs[0]
            run.add_tab()

    doc.add_paragraph()

    # ── 과제 3 개요 ──
    _heading(doc, "과제 3. 문헌 비평과 적용", size=13, bold=True, center=True, sb=10, sa=6)
    for txt in ["1. 임상시험의 결과는 타당한가?",
                "2. 결과는 임상적으로 중요한가?",
                "3. 연구 결과를 실제 환자에게 적용할 수 있는가?"]:
        _para(doc, txt, indent=2.5, sa=3)
    _para(doc, "(별도 워크시트 사용)", indent=7, sa=3)

    _page_break(doc)

    # ══════════════════════════════════════
    # 2페이지: 치료 논문 읽기 워크시트
    # ══════════════════════════════════════
    _heading(doc, "치료에 관한 개별 논문 읽기", size=14, bold=True, center=True, sb=0, sa=12)

    # ── 1) Internal Validity ──
    _para(doc, "1) 임상시험의 결과는 타당한가?", bold=True, size=11, sb=4, sa=4)

    iv_qs = [
        ("치료에 대한 환자 배정은 무작위적인가?",
         r"치료에 대한 환자 배정"),
        ("환자에 대한 추적관찰은 충분히 완전한가?",
         r"추적관찰은 충분히 완전"),
        ("환자에 대한 추적관찰은 질병결과를 관찰하기에 충분히 긴가?",
         r"충분히 긴가"),
        ("환자는 모두 애초에 무작위 배정된 군에 따라 분석되었는가?\n(Intention to treat analysis)",
         r"Intention to treat|ITT"),
        ("이중 맹검법이 시행되는가?",
         r"이중 맹검"),
        ("각 군은 시험의 대상이 되는 치료법 외에는 모든 측면에서 동일하게 취급되었는가?",
         r"동일하게 취급"),
        ("각 군은 임상시험의 시작단계에서 유사하였는가?",
         r"시작단계에서 유사"),
    ]

    iv_tbl = doc.add_table(rows=len(iv_qs), cols=2)
    _borders(iv_tbl)
    for i, (q_label, q_pat) in enumerate(iv_qs):
        ans, ev = _iv(ai_text, q_pat)
        row = iv_tbl.rows[i]
        _cell(row.cells[0], q_label, size=10)
        row.cells[0].width = Cm(11)
        row.cells[1].width = Cm(4)

        # 오른쪽 셀: 예 / 아니오 (답에 따라 강조) + 근거
        ans_color = (0, 112, 192) if ans == "예" else (192, 0, 0) if ans == "아니오" else None
        lines = []
        if ans == "예":
            lines.append(("예  /  아니오", True, (0, 112, 192)))
        elif ans == "아니오":
            lines.append(("예  /  아니오", True, (192, 0, 0)))
        else:
            lines.append(("예  /  아니오", False, None))
        if ev:
            lines.append((f"[근거] {ev[:120]}", False, (80, 80, 80)))
        _cell_multiline(row.cells[1], lines, size=9.5, center=False)
        _set_row_height(row, 1.1)

    doc.add_paragraph()

    # ── 2) 임상적 중요성 ──
    _para(doc, "2) 해당 연구의 결과가 임상적으로 중요한가?", bold=True, size=11, sb=6, sa=4)
    _para(doc, "1. 치료의 효과는 어느 정도인가? (치료군 vs. 대조군)", bold=True, size=10.5, sa=4)

    imp_sec = re.search(r'###\s*2[.\s].*?\n(.*?)(?=\n###|\n##|\Z)', ai_text, re.DOTALL)
    imp_text = imp_sec.group(1) if imp_sec else ai_text

    out_label = _get(imp_text, r"Outcome\s*항목") or _get(ai_text, r"Outcome\s*항목")
    cer_val   = _get(imp_text, r"CER\s*\(Control") or _get(ai_text, r"CER\s*\(Control")
    eer_val   = _get(imp_text, r"EER\s*\(Experimental") or _get(ai_text, r"EER\s*\(Experimental")
    rrr_val   = _get(imp_text, r"RRR") or _get(ai_text, r"RRR")
    arr_val   = _get(imp_text, r"ARR") or _get(ai_text, r"ARR")
    nnt_val   = _get(imp_text, r"NNT\s*\(Number") or _get(ai_text, r"NNT\s*\(Number")

    # 통계 표: Outcome | CER | EER | RRR | ARR | NNT
    st_tbl = doc.add_table(rows=3, cols=6)
    _borders(st_tbl)
    headers = ["Outcome 항목", "CER", "EER",
               "Relative Risk\nReduction\n(RRR)\n|CER-EER|/CER",
               "Absolute Risk\nReduction\n(ARR)\n|CER-EER|",
               "Number Needed\nto Treat\n(NNT)\n1/ARR"]
    values  = [out_label, cer_val, eer_val, rrr_val, arr_val, nnt_val]
    widths  = [3.5, 1.7, 1.7, 3.0, 3.0, 3.0]
    for j, (h, v, w) in enumerate(zip(headers, values, widths)):
        _cell(st_tbl.rows[0].cells[j], h, bold=True, size=9, center=True)
        _cell(st_tbl.rows[1].cells[j], v, size=10, center=True)
        _cell(st_tbl.rows[2].cells[j], "",size=10, center=True)
        st_tbl.rows[0].cells[j].width = Cm(w)
        st_tbl.rows[1].cells[j].width = Cm(w)
        _set_row_height(st_tbl.rows[0], 1.6)
        _set_row_height(st_tbl.rows[1], 1.0)
        _set_row_height(st_tbl.rows[2], 1.0)

    doc.add_paragraph()

    # 95% CI
    _para(doc, "2. 치료효과에 대한 추정은 얼마나 정밀한가?", bold=True, size=10.5, sa=3)
    ci_m = re.search(r'95.*?CI.*?ARR.*?=\s*(.+?)(?=\n\*|\n##|\Z)', ai_text, re.DOTALL|re.IGNORECASE)
    ci_val = _clean(ci_m.group(1)) if ci_m else ""
    se_m = re.search(r'S\.?E\.?\s*=\s*(.+?)(?=\n|\Z)', ai_text, re.DOTALL|re.IGNORECASE)
    se_val = _clean(se_m.group(1)) if se_m else ""
    _para(doc, f"95% CI of ARR = ARR ± 1.96 × S.E.  →  {ci_val}", sa=2)
    _para(doc, f"S.E = √(CER(1-CER)/N₁ + EER(1-EER)/N₂)  →  {se_val}", sa=2)

    _page_break(doc)

    # ══════════════════════════════════════
    # 3페이지: 부작용 + 적용 가능성
    # ══════════════════════════════════════

    # ── 3. 부작용 표 ──
    _para(doc, "3. 치료에 따른 부작용은 어느 정도인가?", bold=True, size=11, sb=0, sa=4)

    adv_sec = re.search(r'###\s*3[.\s].*?\n(.*?)(?=\n###|\n##|\Z)', ai_text, re.DOTALL)
    adv_text = adv_sec.group(1) if adv_sec else ai_text

    adv_item = _get(adv_text, r"주요 부작용 항목") or _get(ai_text, r"주요 부작용 항목")
    adv_cer  = _get(adv_text, r"CER") or ""
    adv_eer  = _get(adv_text, r"EER") or ""
    rri_val  = _get(adv_text, r"RRI") or _get(ai_text, r"RRI")
    ari_val  = _get(adv_text, r"ARI") or _get(ai_text, r"ARI")
    nnh_val  = _get(adv_text, r"NNH\s*\(Number") or _get(ai_text, r"NNH\s*\(Number")

    adv_tbl = doc.add_table(rows=3, cols=6)
    _borders(adv_tbl)
    adv_hdrs = ["주요 부작용", "CER", "EER",
                "Relative Risk\nIncrease\n(RRI)\n|CER-EER|/CER",
                "Absolute Risk\nIncrease\n(ARI)\n|CER-EER|",
                "Number Needed\nto Harm\n(NNH)\n1/ARI"]
    adv_vals = [adv_item, adv_cer, adv_eer, rri_val, ari_val, nnh_val]
    for j, (h, v) in enumerate(zip(adv_hdrs, adv_vals)):
        _cell(adv_tbl.rows[0].cells[j], h, bold=True, size=9, center=True)
        _cell(adv_tbl.rows[1].cells[j], v, size=10, center=True)
        _cell(adv_tbl.rows[2].cells[j], "", size=10, center=True)
        _set_row_height(adv_tbl.rows[0], 1.6)
        _set_row_height(adv_tbl.rows[1], 1.0)
        _set_row_height(adv_tbl.rows[2], 1.0)

    doc.add_paragraph()

    # ── 3) 실제 환자 적용 ──
    _para(doc, "3) 해당연구의 결과를 실제 환자에게 적용할 수 있는가?", bold=True, size=11, sb=6, sa=4)

    app_sec = re.search(r'###\s*4[.\s].*?\n(.*?)(?=\n###|\n##|\Z)', ai_text, re.DOTALL)
    app_text = app_sec.group(1) if app_sec else ai_text

    # 예/아니오 두 질문
    q1_m = re.search(r'실제환자와 연구대상.+?(\[예\]|\[아니오\]|예|아니오).+?이유:\s*(.+?)(?=\n\*|\n##|\Z)', ai_text, re.DOTALL)
    q2_m = re.search(r'다른 치료법은 없는가.+?(\[예\]|\[아니오\]|예|아니오).+?이유:\s*(.+?)(?=\n\*|\n##|\Z)', ai_text, re.DOTALL)
    q1_ans = q1_m.group(1).replace('[','').replace(']','') if q1_m else ""
    q1_ev  = _clean(q1_m.group(2) or "") if q1_m else ""
    q2_ans = q2_m.group(1).replace('[','').replace(']','') if q2_m else ""
    q2_ev  = _clean(q2_m.group(2) or "") if q2_m else ""

    # f-factor 값
    ft_m  = re.search(r'f.*?treatment.*?=\s*([0-9.]+)', ai_text, re.IGNORECASE)
    fa_m  = re.search(r'f.*?adverse.*?=\s*([0-9.]+)', ai_text, re.IGNORECASE)
    nnt_p = re.search(r'NNT.*?patient.*?=\s*(.+?)(?=\n|\Z)', ai_text, re.IGNORECASE)
    nnh_p = re.search(r'NNH.*?patient.*?=\s*(.+?)(?=\n|\Z)', ai_text, re.IGNORECASE)
    f_t   = ft_m.group(1) if ft_m else ""
    f_a   = fa_m.group(1) if fa_m else ""
    nnt_pt= _clean(nnt_p.group(1)) if nnt_p else ""
    nnh_pt= _clean(nnh_p.group(1)) if nnh_p else ""

    # s factor, LHH
    s_m   = re.search(r's\s*(?:factor)?\s*=\s*([0-9.]+)', ai_text, re.IGNORECASE)
    lhh_m = re.search(r'LHH\s*=\s*(.+?)(?=\n##|\n###|\Z)', ai_text, re.DOTALL|re.IGNORECASE)
    s_val  = s_m.group(1) if s_m else ""
    lhh_val= _clean(lhh_m.group(1)) if lhh_m else ""

    # 결론
    conc_m = re.search(r'결론[:\s]*(.+?)(?=\n##|\n###|\Z)', ai_text, re.DOTALL)
    conc_val = _clean(conc_m.group(1)) if conc_m else ""

    # 적용 가능성 표 (PDF 3페이지 구조)
    app_tbl = doc.add_table(rows=8, cols=2)
    _borders(app_tbl)
    app_tbl.columns[0].width = Cm(9)
    app_tbl.columns[1].width = Cm(6)

    def _app_row(r, left, right, lb=False, rb=False):
        _cell(app_tbl.rows[r].cells[0], left,  size=10, bold=lb)
        _cell(app_tbl.rows[r].cells[1], right, size=10, bold=rb)

    _app_row(0,
        "제시된 연구결과를 적용하지 못할 정도로 실제환자와 연구대상 간에 차이가 존재하는가?",
        f"{q1_ans}\n\n{q1_ev[:100]}" if q1_ev else q1_ans)
    _app_row(1, "다른 치료법은 없는가?",
        f"{q2_ans}\n\n{q2_ev[:100]}" if q2_ev else q2_ans)
    _app_row(2,
        "해당 치료법 적용시 실제 환자에서 기대되는 잠재적인 편익과 손실은 무엇인가?\n(NNT와 NNH를 구하여 평가한다.)",
        "")
    _app_row(3,
        "① 'f'를 정한다.",
        f"f for treatment = {f_t}\nf for adverse effect = {f_a}")
    _app_row(4,
        "② 앞서 정한 f를 이용하여 실제환자의 NNT와 NNH를 구한다.",
        f"실제환자의 NNT = NNT/f = {nnt_pt}\n실제환자의 NNH = NNH/f = {nnh_pt}")
    _app_row(5,
        "치료법 자체와 질병치료결과가 실제 환자에게 가지는 가치와 기대는 무엇인가?",
        "")
    _app_row(6,
        "① 's' factor를 구한다. (주어진 scale을 이용할 것)\n② 앞서 구한 f와 s를 이용하여 LHH를 구한다.",
        f"LHH = [(1/NNT)×f×s] vs. [(1/NNH)×f]\n{lhh_val[:150]}")
    for r in range(8):
        _set_row_height(app_tbl.rows[r], 1.2)

    _page_break(doc)

    # ══════════════════════════════════════
    # 4페이지: Rating scale + 결론
    # ══════════════════════════════════════
    _para(doc, "**환자의 질병결과에 대한 가치를 평가하는 rating scale", bold=True, size=11, sb=0, sa=10)

    # Rating scale 1
    _para(doc, "1) 치료를 하지 않았을 때 예상되는 질병 결과", bold=False, size=10.5, sa=4)
    scale_tbl1 = doc.add_table(rows=2, cols=5)
    _borders(scale_tbl1)
    for j, (val, lbl) in enumerate(zip(["0", "0.1", "0.5", "", "1.0"],
                                       ["Outcome","","","","완치(건강)"])):
        _cell(scale_tbl1.rows[0].cells[j], val, center=True, size=10)
        _cell(scale_tbl1.rows[1].cells[j], lbl, center=(j not in [0,4]), size=10)
    scale_tbl1.columns[0].width = Cm(1.5)
    scale_tbl1.columns[1].width = Cm(1.5)
    scale_tbl1.columns[2].width = Cm(4.0)
    scale_tbl1.columns[3].width = Cm(4.0)
    scale_tbl1.columns[4].width = Cm(2.0)

    doc.add_paragraph()

    # Rating scale 2
    _para(doc, "2) 치료에 대한 부작용", bold=False, size=10.5, sa=4)
    scale_tbl2 = doc.add_table(rows=2, cols=5)
    _borders(scale_tbl2)
    for j, (val, lbl) in enumerate(zip(["0", "0.1", "0.5", "", "1.0"],
                                       ["수용 못함","","","","수용 함"])):
        _cell(scale_tbl2.rows[0].cells[j], val, center=True, size=10)
        _cell(scale_tbl2.rows[1].cells[j], lbl, center=(j not in [0,4]), size=10)
    scale_tbl2.columns[0].width = Cm(1.5)
    scale_tbl2.columns[1].width = Cm(1.5)
    scale_tbl2.columns[2].width = Cm(4.0)
    scale_tbl2.columns[3].width = Cm(4.0)
    scale_tbl2.columns[4].width = Cm(2.0)

    doc.add_paragraph()

    # s factor 분수 표시
    _para(doc, f"2)값       {s_val}", size=11, sb=6, sa=0)
    _para(doc, "──────", size=11, sb=0, sa=0)
    _para(doc, "1)값", size=11, sb=0, sa=10)

    # 결론
    p_conc = doc.add_paragraph()
    p_conc.paragraph_format.space_before = Pt(10)
    r1 = p_conc.add_run("결론 :  이제까지 시행한 분석에 근거하여 당신은 환자에게 위의 치료를 하도록 권고할 것인가?   ")
    _font(r1, size=10.5)
    conc_yn = "예" if "예" in conc_val[:10] else "아니오" if "아니오" in conc_val[:15] else ""
    r2 = p_conc.add_run(f"예  /  아니오")
    color = (0,112,192) if conc_yn=="예" else (192,0,0) if conc_yn=="아니오" else None
    _font(r2, size=10.5, bold=True, color=color)
    if conc_val:
        _para(doc, f"(소견) {conc_val[:300]}", size=10, sb=4, sa=4)

    _page_break(doc)

    # ══════════════════════════════════════
    # 5페이지: DISCUSSION
    # ══════════════════════════════════════
    _heading(doc, "DISCUSSION", size=14, bold=True, center=True, sb=0, sa=20)

    disc_m = re.search(r'##\s*DISCUSSION\s*\n(.*)', ai_text, re.DOTALL|re.IGNORECASE)
    disc_full = disc_m.group(1).strip() if disc_m else ""

    item1_m = re.search(r'^\s*1[.\)]\s*(.+?)(?=^\s*2[.\)])', disc_full, re.DOTALL|re.MULTILINE)
    item2_m = re.search(r'^\s*2[.\)]\s*(.+?)(?=^\s*3[.\)]|\Z)', disc_full, re.DOTALL|re.MULTILINE)
    item3_m = re.search(r'^\s*3[.\)]\s*(.+)', disc_full, re.DOTALL|re.MULTILINE)

    disc1 = _clean(item1_m.group(1)) if item1_m else ""
    disc2 = _clean(item2_m.group(1)) if item2_m else ""
    disc3 = ""  # EBM 수업 소감은 AI가 생성하지 않음

    _para(doc, "1. 앞에서 분석한 논문이 가지고 있는 문제점이나 제한점에는 어떠한 것들이 있겠는가?",
          bold=False, size=10.5, sb=4, sa=6)
    _para(doc, disc1 or "", size=10.5, indent=0.5, sa=20)

    _para(doc, "2. 이러한 제한점을 보완하기 위하여 어떤 근거가 더 필요하겠는가?",
          bold=False, size=10.5, sb=4, sa=6)
    _para(doc, disc2 or "", size=10.5, indent=0.5, sa=20)

    _para(doc, "3. ebm 수업에 대한 전반적인 소감?",
          bold=False, size=10.5, sb=4, sa=6)
    _para(doc, "", size=10.5, sa=20)

    # 저장
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()
